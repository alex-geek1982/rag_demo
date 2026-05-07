"""
RAG Engine - Document parsers
"""
from abc import ABC, abstractmethod
from typing import List, Optional, Dict, Any
from pathlib import Path
import logging
from ..types import Document, ContentBlock, ContentType, ModalityType
from ..i18n import get_i18n


logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """Base parser interface"""
    
    def __init__(self, doc_id: str, doc_title: str, language: str = "en"):
        """
        Initialize parser
        
        Args:
            doc_id: Document ID
            doc_title: Document title
            language: Document language
        """
        self.doc_id = doc_id
        self.doc_title = doc_title
        self.language = language
        self.i18n = get_i18n()
    
    @abstractmethod
    def parse(self, file_path: str) -> Document:
        """
        Parse document
        
        Args:
            file_path: Path to document file
        
        Returns:
            Parsed document
        """
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """
        Check if parser supports file type
        
        Args:
            file_path: Path to file
        
        Returns:
            True if supported
        """
        pass


class TextParser(BaseParser):
    """Plain text file parser"""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is plain text"""
        return file_path.lower().endswith(('.txt', '.md', '.rst'))
    
    def parse(self, file_path: str) -> Document:
        """Parse text file"""
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document(
                id=self.doc_id,
                title=self.doc_title,
                source_path=str(path),
                language=self.language
            )
            
            # Create single content block for text file
            block = ContentBlock(
                id=f"{self.doc_id}_text_0",
                type=ContentType.TEXT,
                content=content,
                modality=ModalityType.TEXT,
                language=self.language
            )
            doc.add_content_block(block)
            
            logger.info(f"Successfully parsed text file: {file_path}")
            return doc
        except Exception as e:
            logger.error(f"Parsing failed: {str(e)}")
            raise


class PDFParser(BaseParser):
    """PDF file parser with advanced layout understanding"""
    
    def __init__(
        self,
        doc_id: str,
        doc_title: str,
        language: str = "en",
        use_advanced_layout: bool = True,
        extract_images: bool = True,
        extract_tables: bool = True,
        extract_text: bool = True,
        use_vision_api: bool = True,
        vision_api_key: Optional[str] = None,
        vision_base_url: Optional[str] = None,
        vision_model: Optional[str] = None,
        vision_provider: str = "openai",
        vision_azure_endpoint: Optional[str] = None,
        vision_azure_api_version: Optional[str] = None,
        vision_azure_deployment: Optional[str] = None,
        context_window_pixels: int = 200,
        min_image_area: int = 1000,
        max_surrounding_text_chars: int = 2000,
        filter_header_footer: bool = True,
        header_margin_ratio: float = 0.08,
        footer_margin_ratio: float = 0.08,
        header_footer_min_repeat_pages: int = 2,
    ):
        """
        Initialize PDF parser
        
        Args:
            doc_id: Document ID
            doc_title: Document title
            language: Document language
            use_advanced_layout: Use advanced PDF processing with layout understanding
            extract_images: Extract and describe images from PDF
            extract_tables: Extract structured tables
            use_vision_api: Use Vision API for image descriptions
        """
        super().__init__(doc_id, doc_title, language)
        self.use_advanced_layout = use_advanced_layout
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.extract_text = extract_text
        self.use_vision_api = use_vision_api
        self.vision_api_key = vision_api_key
        self.vision_base_url = vision_base_url
        self.vision_model = vision_model
        self.vision_provider = vision_provider
        self.vision_azure_endpoint = vision_azure_endpoint
        self.vision_azure_api_version = vision_azure_api_version
        self.vision_azure_deployment = vision_azure_deployment
        self.context_window_pixels = context_window_pixels
        self.min_image_area = min_image_area
        self.max_surrounding_text_chars = max_surrounding_text_chars
        self.filter_header_footer = filter_header_footer
        self.header_margin_ratio = header_margin_ratio
        self.footer_margin_ratio = footer_margin_ratio
        self.header_footer_min_repeat_pages = header_footer_min_repeat_pages
        
        if use_advanced_layout:
            self._init_advanced_processor()
    
    def _init_advanced_processor(self) -> None:
        """Initialize advanced PDF processor"""
        try:
            from .pdf_advanced import AdvancedPDFProcessor
            self.advanced_processor = AdvancedPDFProcessor(
                extract_images=self.extract_images,
                extract_tables=self.extract_tables,
                extract_text=self.extract_text,
                use_vision_api=self.use_vision_api,
                vision_api_key=self.vision_api_key,
                vision_base_url=self.vision_base_url,
                vision_model=self.vision_model,
                vision_provider=self.vision_provider,
                vision_azure_endpoint=self.vision_azure_endpoint,
                vision_azure_api_version=self.vision_azure_api_version,
                vision_azure_deployment=self.vision_azure_deployment,
                context_window_pixels=self.context_window_pixels,
                min_image_area=self.min_image_area,
                max_surrounding_text_chars=self.max_surrounding_text_chars,
                filter_header_footer=self.filter_header_footer,
                header_margin_ratio=self.header_margin_ratio,
                footer_margin_ratio=self.footer_margin_ratio,
                header_footer_min_repeat_pages=self.header_footer_min_repeat_pages,
            )
        except ImportError as e:
            logger.warning(f"Advanced PDF processor not available: {e}")
            self.advanced_processor = None
    
    def supports(self, file_path: str) -> bool:
        """Check if file is PDF"""
        return file_path.lower().endswith('.pdf')
    
    def parse(self, file_path: str) -> Document:
        """Parse PDF file with layout understanding"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Try advanced processing first
        if self.use_advanced_layout and self.advanced_processor:
            try:
                logger.info(f"Using advanced PDF processing for {file_path}")
                return self.advanced_processor.process_pdf(
                    file_path=str(path),
                    doc_id=self.doc_id,
                    doc_title=self.doc_title,
                    language=self.language
                )
            except Exception as e:
                logger.warning(f"Advanced processing failed, falling back to basic: {e}")
        
        # Fallback to basic PDF processing
        return self._parse_pdf_basic(str(path))
    
    def _parse_pdf_basic(self, file_path: str) -> Document:
        """Basic PDF parsing without advanced features"""
        try:
            import PyPDF2
            
            path = Path(file_path)
            doc = Document(
                id=self.doc_id,
                title=self.doc_title,
                source_path=str(path),
                language=self.language
            )
            
            with open(path, 'rb') as f:
                pdfReader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdfReader.pages)):
                    page = pdfReader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        block = ContentBlock(
                            id=f"{self.doc_id}_pdf_page_{page_num}",
                            type=ContentType.TEXT,
                            content=text,
                            modality=ModalityType.TEXT,
                            page_num=page_num,
                            language=self.language
                        )
                        doc.add_content_block(block)
            
            logger.info(f"Successfully parsed PDF file (basic mode): {file_path}")
            return doc
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            raise


class DocxParser(BaseParser):
    """DOCX file parser"""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is DOCX"""
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def parse(self, file_path: str) -> Document:
        """Parse DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            docx_doc = DocxDocument(path)
            doc = Document(
                id=self.doc_id,
                title=self.doc_title,
                source_path=str(path),
                language=self.language
            )
            
            block_idx = 0
            for para_idx, para in enumerate(docx_doc.paragraphs):
                if para.text.strip():
                    block = ContentBlock(
                        id=f"{self.doc_id}_docx_para_{block_idx}",
                        type=ContentType.TEXT,
                        content=para.text,
                        modality=ModalityType.TEXT,
                        metadata={"paragraph_index": para_idx},
                        language=self.language
                    )
                    doc.add_content_block(block)
                    block_idx += 1
            
            # Extract table content
            for table_idx, table in enumerate(docx_doc.tables):
                table_content = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_content.append("|".join(row_text))
                
                if table_content:
                    table_text = "\n".join(table_content)
                    block = ContentBlock(
                        id=f"{self.doc_id}_docx_table_{table_idx}",
                        type=ContentType.TABLE,
                        content=table_text,
                        modality=ModalityType.STRUCTURED,
                        metadata={"table_index": table_idx},
                        language=self.language
                    )
                    doc.add_content_block(block)
            
            logger.info(f"Successfully parsed DOCX file: {file_path}")
            return doc
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Parsing failed: {str(e)}")
            raise


class ExcelParser(BaseParser):
    """Excel file parser (XLSX, XLS)"""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is Excel"""
        return file_path.lower().endswith(('.xlsx', '.xls'))
    
    def parse(self, file_path: str) -> Document:
        """Parse Excel file"""
        try:
            import pandas as pd
            
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            doc = Document(
                id=self.doc_id,
                title=self.doc_title,
                source_path=str(path),
                language=self.language
            )
            
            # Read all sheets
            xls_file = pd.ExcelFile(path)
            for sheet_idx, sheet_name in enumerate(xls_file.sheet_names):
                df = pd.read_excel(path, sheet_name=sheet_name)
                
                # Convert DataFrame to markdown table format
                table_content = df.to_markdown(index=False)
                
                block = ContentBlock(
                    id=f"{self.doc_id}_excel_sheet_{sheet_idx}",
                    type=ContentType.TABLE,
                    content=table_content,
                    modality=ModalityType.STRUCTURED,
                    metadata={"sheet_name": sheet_name, "sheet_index": sheet_idx},
                    language=self.language
                )
                doc.add_content_block(block)
            
            logger.info(f"Successfully parsed Excel file: {file_path}")
            return doc
        except ImportError:
            logger.error("pandas not installed. Install with: pip install pandas openpyxl")
            raise
        except Exception as e:
            logger.error(f"Parsing failed: {str(e)}")
            raise


class ImageParser(BaseParser):
    """Image file parser"""
    
    def supports(self, file_path: str) -> bool:
        """Check if file is image"""
        return file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff'))
    
    def parse(self, file_path: str) -> Document:
        """Parse image file"""
        try:
            from PIL import Image
            
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Verify it's a valid image
            img = Image.open(path)
            img_format = img.format
            
            doc = Document(
                id=self.doc_id,
                title=self.doc_title,
                source_path=str(path),
                language=self.language
            )
            
            # Create image content block
            block = ContentBlock(
                id=f"{self.doc_id}_image_0",
                type=ContentType.IMAGE,
                content=str(path),  # Store image path
                modality=ModalityType.VISUAL,
                metadata={
                    "format": img_format,
                    "size": f"{img.width}x{img.height}",
                    "mode": img.mode
                },
                language=self.language
            )
            doc.add_content_block(block)
            
            logger.info(f"Successfully parsed image file: {file_path}")
            return doc
        except ImportError:
            logger.error("Pillow not installed. Install with: pip install Pillow")
            raise
        except Exception as e:
            logger.error(f"Parsing failed: {str(e)}")
            raise


class ParserFactory:
    """Factory for creating appropriate parser"""
    
    PARSERS = [
        PDFParser,
        DocxParser,
        ExcelParser,
        ImageParser,
        TextParser,  # Default - should be last
    ]
    
    @classmethod
    def get_parser(
        cls,
        file_path: str,
        doc_id: str,
        doc_title: str,
        language: str = "en",
        pdf_config: Optional[Any] = None,
        vision_config: Optional[Any] = None,
    ) -> BaseParser:
        """
        Get appropriate parser for file
        
        Args:
            file_path: Path to file
            doc_id: Document ID
            doc_title: Document title
            language: Document language
            pdf_config: Optional PDF processing config
            vision_config: Optional vision model config
        
        Returns:
            Appropriate parser instance
        """
        for parser_class in cls.PARSERS:
            if parser_class is PDFParser:
                # Resolve vision configuration from both pdf_config and vision_config
                vision_api_key = (
                    getattr(vision_config, "api_key", None)  # Priority 1: vision_config.api_key
                    or getattr(pdf_config, "vision_api_key", None)  # Priority 2: pdf_config.vision_api_key
                )
                vision_base_url = (
                    getattr(pdf_config, "vision_base_url", None)  # Priority 1: pdf_config.vision_base_url
                    or getattr(vision_config, "base_url", None)  # Priority 2: vision_config.base_url
                )
                vision_model = (
                    getattr(pdf_config, "vision_model", None)  # Priority 1: pdf_config.vision_model
                    or getattr(vision_config, "model", None)  # Priority 2: vision_config.model
                )
                vision_provider = (
                    getattr(vision_config, "provider", None)  # Priority 1: vision_config.provider
                    or getattr(pdf_config, "vision_provider", None)  # Priority 2: pdf_config.vision_provider
                    or "openai"  # Default
                )
                vision_azure_endpoint = (
                    getattr(pdf_config, "vision_azure_endpoint", None)
                    or getattr(vision_config, "azure_endpoint", None)
                )
                vision_azure_api_version = (
                    getattr(pdf_config, "vision_azure_api_version", None)
                    or getattr(vision_config, "azure_api_version", None)
                )
                vision_azure_deployment = (
                    getattr(pdf_config, "vision_azure_deployment", None)
                    or getattr(vision_config, "azure_deployment", None)
                )

                parser = parser_class(
                    doc_id,
                    doc_title,
                    language,
                    use_advanced_layout=getattr(pdf_config, "use_advanced_layout", True),
                    extract_images=getattr(pdf_config, "extract_images", True),
                    extract_tables=getattr(pdf_config, "extract_tables", True),
                    extract_text=getattr(pdf_config, "extract_text", True),
                    use_vision_api=getattr(pdf_config, "use_vision_api", True),
                    vision_api_key=vision_api_key,
                    vision_base_url=vision_base_url,
                    vision_model=vision_model,
                    vision_provider=vision_provider,
                    vision_azure_endpoint=vision_azure_endpoint,
                    vision_azure_api_version=vision_azure_api_version,
                    vision_azure_deployment=vision_azure_deployment,
                    context_window_pixels=getattr(pdf_config, "context_window_pixels", 200),
                    min_image_area=getattr(pdf_config, "min_image_area", 1000),
                    max_surrounding_text_chars=getattr(pdf_config, "max_surrounding_text_chars", 2000),
                    filter_header_footer=getattr(pdf_config, "filter_header_footer", True),
                    header_margin_ratio=getattr(pdf_config, "header_margin_ratio", 0.08),
                    footer_margin_ratio=getattr(pdf_config, "footer_margin_ratio", 0.08),
                    header_footer_min_repeat_pages=getattr(pdf_config, "header_footer_min_repeat_pages", 2),
                )
            else:
                parser = parser_class(doc_id, doc_title, language)

            if parser.supports(file_path):
                return parser
        
        raise ValueError(f"No parser found for file: {file_path}")

